import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from models import ChapterContent, AgentStatusEnum
from typing import List
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

class DocumentFormatter:
    def __init__(self):
        self.name = "DocumentFormatter"
        
    async def update_progress(self, db, project_id: str, status: AgentStatusEnum, 
                            progress: float, task: str = None):
        """Update progress in database"""
        await db.agent_progress.update_one(
            {"project_id": project_id, "agent_name": self.name},
            {
                "$set": {
                    "status": status.value,
                    "progress_percentage": progress,
                    "current_task": task,
                    "updated_at": datetime.utcnow()
                }
            },
            upsert=True
        )
    
    async def create_kdp_document(self, title: str, chapters: List[ChapterContent], 
                                project_id: str) -> str:
        """Create a KDP-ready Word document"""
        try:
            # Create new document
            doc = Document()
            
            # Set up KDP formatting (8.5" x 11")
            self._setup_kdp_formatting(doc)
            
            # Add title page
            self._add_title_page(doc, title)
            
            # Add table of contents
            self._add_table_of_contents(doc, chapters)
            
            # Add chapters
            for i, chapter in enumerate(chapters):
                self._add_chapter(doc, chapter, i == 0)
            
            # Save document
            filename = f"{title.replace(' ', '_').replace('/', '_')}_{project_id}.docx"
            filepath = f"/app/backend/generated_books/{filename}"
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            doc.save(filepath)
            
            logger.info(f"KDP document created: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error creating KDP document: {str(e)}")
            raise
    
    def _setup_kdp_formatting(self, doc: Document):
        """Set up KDP-specific formatting for 8.5" x 11" page"""
        # Get the section
        section = doc.sections[0]
        
        # Set page size to 8.5" x 11" (KDP standard)
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        
        # Set margins for KDP (minimum 0.25" on all sides, but 0.5" is safer)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.625)  # Slightly larger for binding
        section.right_margin = Inches(0.5)
        
        # Set header/footer margins
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        
        # Set paragraph spacing
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.15
    
    def _add_title_page(self, doc: Document, title: str):
        """Add a professional title page"""
        # Add title
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(24)
        title_run.bold = True
        
        # Add some space
        for _ in range(8):
            doc.add_paragraph()
        
        # Add author placeholder
        author_paragraph = doc.add_paragraph()
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_paragraph.add_run("by [Author Name]")
        author_run.font.name = 'Times New Roman'
        author_run.font.size = Pt(14)
        
        # Add page break
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc: Document, chapters: List[ChapterContent]):
        """Add table of contents"""
        # TOC Title
        toc_title = doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_title.add_run("Table of Contents")
        toc_run.font.name = 'Times New Roman'
        toc_run.font.size = Pt(16)
        toc_run.bold = True
        
        doc.add_paragraph()
        
        # Add chapters to TOC
        for chapter in chapters:
            toc_entry = doc.add_paragraph()
            toc_entry.add_run(f"Chapter {chapter.chapter_number}: {chapter.title}")
            toc_entry.add_run("\t")
            toc_entry.add_run(f"{chapter.chapter_number + 2}")  # Page numbers (approximate)
        
        # Add page break
        doc.add_page_break()
    
    def _add_chapter(self, doc: Document, chapter: ChapterContent, is_first: bool = False):
        """Add a chapter to the document"""
        if not is_first:
            # Add page break before each chapter (except first)
            doc.add_page_break()
        
        # Chapter title
        chapter_title = doc.add_paragraph()
        chapter_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Chapter number
        chapter_num_run = chapter_title.add_run(f"Chapter {chapter.chapter_number}")
        chapter_num_run.font.name = 'Times New Roman'
        chapter_num_run.font.size = Pt(14)
        chapter_num_run.bold = True
        
        # Chapter title
        chapter_title.add_run("\n")
        title_run = chapter_title.add_run(chapter.title)
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(16)
        title_run.bold = True
        
        # Add space before content
        doc.add_paragraph()
        
        # Add chapter content
        self._add_formatted_content(doc, chapter.content)
    
    def _add_formatted_content(self, doc: Document, content: str):
        """Add formatted content with proper paragraph breaks"""
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                # Handle dialogue and regular paragraphs
                para = doc.add_paragraph()
                
                # Check if it's dialogue (starts with quote)
                if para_text.strip().startswith('"'):
                    # Indent dialogue slightly
                    para.paragraph_format.first_line_indent = Inches(0.25)
                else:
                    # Regular paragraph indent
                    para.paragraph_format.first_line_indent = Inches(0.5)
                
                # Add the text
                run = para.add_run(para_text.strip())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                
                # Justify text
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _add_headers_footers(self, doc: Document, title: str):
        """Add headers and footers for professional look"""
        section = doc.sections[0]
        
        # Header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run(title)
        header_run.font.name = 'Times New Roman'
        header_run.font.size = Pt(9)
        header_run.italic = True
        
        # Footer with page numbers
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        footer_para._p.append(fldChar1)
        footer_para._p.append(instrText)
        footer_para._p.append(fldChar2)
    
    def get_file_size_mb(self, filepath: str) -> float:
        """Get file size in MB"""
        try:
            size_bytes = os.path.getsize(filepath)
            return size_bytes / (1024 * 1024)
        except:
            return 0.0
    
    def validate_kdp_requirements(self, filepath: str) -> Dict[str, Any]:
        """Validate that the document meets KDP requirements"""
        validation_results = {
            "valid": True,
            "warnings": [],
            "errors": []
        }
        
        try:
            # Check file size (KDP limit is around 650MB, but much smaller is better)
            file_size_mb = self.get_file_size_mb(filepath)
            if file_size_mb > 50:  # 50MB warning threshold
                validation_results["warnings"].append(f"Large file size: {file_size_mb:.2f}MB")
            
            # Check if file exists
            if not os.path.exists(filepath):
                validation_results["errors"].append("File does not exist")
                validation_results["valid"] = False
            
            # You could add more validation here:
            # - Check document structure
            # - Validate formatting
            # - Check for proper margins
            # - Verify fonts are embeddable
            
        except Exception as e:
            validation_results["errors"].append(f"Validation error: {str(e)}")
            validation_results["valid"] = False
        
        return validation_results